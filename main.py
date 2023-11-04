import numpy as np
import matplotlib.pyplot as plt
import cv2
import pandas as pd
from docx import Document
from docx.shared import Inches
from io import BytesIO

# Część pierwsza — wczytywanie obrazów
imgA1 = plt.imread('A1.png')
imgA2 = plt.imread('A2.jpg')
imgA3 = plt.imread('A3.png')
imgA4 = plt.imread('A4.jpg')
imgB1 = plt.imread('B01.png')
imgB2 = plt.imread('B02.jpg')

# print(imgA1.dtype)
# print(imgA1.shape)
# print(np.min(imgA1), np.max(imgA1))
# test = cv2.imread('A1.png')
# print(test)


# # Zadanie Obraz 1
def imgToUint8(img):  #
    print(np.issubdtype(img.dtype, np.integer))  # A2 A4 B2
    print(np.issubdtype(img.dtype, np.unsignedinteger))  # A2 A4 B2
    print(np.issubdtype(img.dtype, np.floating))  # A1 A3 B1
    print('Stare wartości')
    print(img.dtype)
    print(img.shape)
    print(np.min(img), np.max(img))
    newImg = (img * 255.0).astype('uint8')
    print('Nowe wartości')
    print(newImg.dtype)
    print(newImg.shape)
    print(np.min(newImg), np.max(newImg))
    plt.imshow(newImg)
    # plt.show()


def imgToFloat(img):
    print(np.issubdtype(img.dtype, np.integer))  # A2 A4 B2
    print(np.issubdtype(img.dtype, np.unsignedinteger))  # A2 A4 B2
    print(np.issubdtype(img.dtype, np.floating))  # A1 A3 B1
    print('Stare wartości')
    print(img.dtype)
    print(img.shape)
    print(np.min(img), np.max(img))
    newImg = img / 255.0
    print('Nowe wartości')
    print(newImg.dtype)
    print(newImg.shape)
    print(np.min(newImg), np.max(newImg))
    plt.imshow(newImg)
    # plt.show()


# Część druga — wyświetlanie obrazów
def imgToGray1(img):
    R = img[:, :, 0]
    plt.imshow(R, cmap=plt.cm.gray, vmin=0, vmax=255)
    # plt.show()


def imgToGray2(img):
    img_BGR = cv2.imread(img)
    img_gray = cv2.cvtColor(img_BGR, cv2.COLOR_BGR2GRAY)
    plt.imshow(img_gray, cmap='gray')
    # plt.show()

    # r, g, b = img[:, :, 0], img[:, :, 1], img[:, :, 2]
    # # result = 0.2126 * r + 0.7152 * g + 0.0722 * b # inny odcien w materiałach
    # result = 0.2989 * r + 0.5870 * g + 0.1140 * b
    # plt.imshow(result, cmap='gray')
    # plt.show()


imgU = imgB1  # A1 A3 B1
# imgToUint8(imgU)

imgF = imgB2  # A2 A4 B2


# imgToFloat(imgF)

# imgGrey1 = imgB2
# imgGrey2 = 'B02.jpg'
# imgToGray1(imgGrey1)
# imgToGray2(imgGrey2)

# Zadanie Obraz 2
def partThree(img, saveImg):
    plt.figure(figsize=(10, 10))
    fig, axs = plt.subplots(3, 3)
    plt.subplot(3, 3, 1)
    plt.title("Oryginalny")
    if saveImg:
        plt.imsave('200_200_default.png', img)
    plt.imshow(img)

    plt.subplot(3, 3, 2)
    plt.title('Y1')
    r1, g1, b1 = img[:, :, 0], img[:, :, 1], img[:, :, 2]
    result1 = 0.299 * r1 + 0.587 * g1 + 0.114 * b1
    if saveImg:
        plt.imsave('200_200_Y1.png', result1, cmap='gray')
    plt.imshow(result1, cmap='gray')

    plt.subplot(3, 3, 3)
    plt.title('Y2')
    r2, g2, b2 = img[:, :, 0], img[:, :, 1], img[:, :, 2]
    result2 = 0.2126 * r2 + 0.7152 * g2 + 0.0722 * b2
    if saveImg:
        plt.imsave('200_200_Y2.png', result2, cmap='gray')
    plt.imshow(result2, cmap='gray')

    plt.subplot(3, 3, 4)
    plt.title('R')
    R = img[:, :, 0]
    if saveImg:
        plt.imsave('200_200_R.png', R, cmap=plt.cm.gray, vmin=0, vmax=255)
    plt.imshow(R, cmap=plt.cm.gray, vmin=0, vmax=255)

    plt.subplot(3, 3, 5)
    plt.title('G')
    G = img[:, :, 1]
    if saveImg:
        plt.imsave('200_200_G.png', G, cmap=plt.cm.gray, vmin=0, vmax=255)
    plt.imshow(G, cmap=plt.cm.gray, vmin=0, vmax=255)

    plt.subplot(3, 3, 6)
    plt.title('B')
    B = img[:, :, 2]
    if saveImg:
        plt.imsave('200_200_B.png', B, cmap=plt.cm.gray, vmin=0, vmax=255)
    plt.imshow(B, cmap=plt.cm.gray, vmin=0, vmax=255)

    plt.subplot(3, 3, 7)
    plt.title('R')
    imgTmp = img.copy()
    imgTmp[:, :, 1] = 0
    imgTmp[:, :, 2] = 0
    if saveImg:
        plt.imsave('200_200_RR.png', imgTmp)
    plt.imshow(imgTmp)

    plt.subplot(3, 3, 8)
    plt.title('G')
    imgTmp = img.copy()
    imgTmp[:, :, 0] = 0
    imgTmp[:, :, 2] = 0
    if saveImg:
        plt.imsave('200_200_GG.png', imgTmp)
    plt.imshow(imgTmp)

    plt.subplot(3, 3, 9)
    plt.title('B')
    imgTmp = img.copy()
    imgTmp[:, :, 0] = 0
    imgTmp[:, :, 1] = 0
    if saveImg:
        plt.imsave('200_200_BB.png', imgTmp)
    plt.imshow(imgTmp)
    plt.savefig('wykresP3.jpg')
    return fig


imgPartThree = imgB2
partThree(imgPartThree, 0)

# zadanie trzecie

df = pd.DataFrame(data={'Filename': ['B02.jpg'],
                        'Fragments': [[[0, 0, 200, 200], [300, 300, 500, 500], [600, 600, 800, 800]]]
                        })
# print(df)

document = Document()
document.add_heading('Łukasz Kaszewski - kl49864\nSystemy multimedialne - lab02', 0)  # tworzenie nagłówków druga wartość to poziom nagłówka
coordinades = 0
for index, row in df.iterrows():
    img = plt.imread(row['Filename'])
    if row['Fragments'] is not None:
        # mamy nie pustą listę fragmentów

        for f in row['Fragments']:
            fragment = img[f[0]:f[2],f[1]:f[3]].copy()

        ############################################################
            f = partThree(fragment, 0)
        ############################################################

            f.suptitle('Fragment {}'.format(df.loc[0, 'Fragments'][coordinades]))  # Tytuł wykresu
            f.tight_layout(pad=1.5)  # poprawa czytelności
            memfile = BytesIO()  # tworzenie bufora
            f.savefig(memfile)  # z zapis do bufora
            document.add_picture(memfile, width=Inches(6))  # dodanie obrazu z bufora do pliku
            memfile.close()
            coordinades += 1
        ############################################################
        # Tu dodajesz dane tekstowe - wartosci, wyjscie funkcji ect.

        ############################################################

    document.save('kl49864 Łukasz Kaszewski.docx')  # zapis do pliku
